from __future__ import annotations

import json
from pathlib import Path
from typing import Dict

import pandas as pd


def _round_df(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    for c in ["IDI_raw", "EMI_raw", "MTI_raw"]:
        if c in out.columns:
            out[c] = pd.to_numeric(out[c], errors="coerce").round(8)
    for c in ["IDI_percent", "EMI_percent", "MTI_percent"]:
        if c in out.columns:
            out[c] = pd.to_numeric(out[c], errors="coerce").round(6)
    if "EVI_raw" in out.columns:
        out["EVI_raw"] = pd.to_numeric(out["EVI_raw"], errors="coerce").round(2)
    if "EVI_norm" in out.columns:
        out["EVI_norm"] = pd.to_numeric(out["EVI_norm"], errors="coerce").round(4)
    for c in ["IP_context", "IP_abs_context"]:
        if c in out.columns:
            out[c] = pd.to_numeric(out[c], errors="coerce").round(6)
    for c in out.columns:
        if c.endswith("percentile"):
            out[c] = pd.to_numeric(out[c], errors="coerce").round(2)
    return out


def export_all(
    out_dir: Path,
    texts_df: pd.DataFrame,
    contexts_df: pd.DataFrame,
    distributions_df: pd.DataFrame,
    candidate_df: pd.DataFrame,
    verified_df: pd.DataFrame,
    rejected_df: pd.DataFrame,
    change_log_df: pd.DataFrame,
    quality_flags_df: pd.DataFrame,
    report_md: str,
) -> Dict[str, Path]:
    out_dir.mkdir(parents=True, exist_ok=True)

    texts = _round_df(texts_df)
    contexts = _round_df(contexts_df)
    dists = _round_df(distributions_df)

    paths: Dict[str, Path] = {}

    paths["calibration_texts_csv"] = out_dir / "calibration_texts.csv"
    paths["calibration_texts_xlsx"] = out_dir / "calibration_texts.xlsx"
    texts.to_csv(paths["calibration_texts_csv"], index=False)
    with pd.ExcelWriter(paths["calibration_texts_xlsx"], engine="openpyxl") as xw:
        texts.to_excel(xw, index=False, sheet_name="calibration_texts")

    paths["calibration_contexts_csv"] = out_dir / "calibration_contexts.csv"
    paths["calibration_contexts_xlsx"] = out_dir / "calibration_contexts.xlsx"
    contexts.to_csv(paths["calibration_contexts_csv"], index=False)
    with pd.ExcelWriter(paths["calibration_contexts_xlsx"], engine="openpyxl") as xw:
        contexts.to_excel(xw, index=False, sheet_name="calibration_contexts")

    paths["calibration_distributions_csv"] = out_dir / "calibration_distributions.csv"
    paths["calibration_distributions_xlsx"] = out_dir / "calibration_distributions.xlsx"
    paths["calibration_distributions_json"] = out_dir / "calibration_distributions.json"
    dists.to_csv(paths["calibration_distributions_csv"], index=False)
    with pd.ExcelWriter(paths["calibration_distributions_xlsx"], engine="openpyxl") as xw:
        dists.to_excel(xw, index=False, sheet_name="calibration_distributions")
    paths["calibration_distributions_json"].write_text(dists.to_json(orient="records", force_ascii=False, indent=2), encoding="utf-8")

    paths["candidate_terms_csv"] = out_dir / "candidate_terms.csv"
    paths["verified_terms_csv"] = out_dir / "verified_terms.csv"
    paths["rejected_terms_csv"] = out_dir / "rejected_terms.csv"
    paths["dictionary_change_log_csv"] = out_dir / "dictionary_change_log.csv"
    paths["calibration_quality_flags_csv"] = out_dir / "calibration_quality_flags.csv"
    candidate_df.to_csv(paths["candidate_terms_csv"], index=False)
    verified_df.to_csv(paths["verified_terms_csv"], index=False)
    rejected_df.to_csv(paths["rejected_terms_csv"], index=False)
    change_log_df.to_csv(paths["dictionary_change_log_csv"], index=False)
    quality_flags_df.to_csv(paths["calibration_quality_flags_csv"], index=False)

    paths["calibration_report_md"] = out_dir / "calibration_report.md"
    paths["calibration_report_md"].write_text(report_md, encoding="utf-8")
    try:
        from docx import Document  # type: ignore

        doc = Document()
        for line in report_md.splitlines():
            if line.startswith("# "):
                doc.add_heading(line[2:].strip(), level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:].strip(), level=2)
            else:
                doc.add_paragraph(line)
        paths["calibration_report_docx"] = out_dir / "calibration_report.docx"
        doc.save(str(paths["calibration_report_docx"]))
    except Exception:
        pass

    paths["calibration_interpretation_rules_json"] = out_dir / "calibration_interpretation_rules.json"
    rules = {
        "levels": {
            "very_low": "0-20",
            "low": "21-40",
            "medium": "41-60",
            "elevated": "61-80",
            "high": "81-95",
            "extreme": "96-100",
        }
    }
    paths["calibration_interpretation_rules_json"].write_text(json.dumps(rules, ensure_ascii=False, indent=2), encoding="utf-8")

    return paths
