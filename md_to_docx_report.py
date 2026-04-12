#!/usr/bin/env python3
"""Convert markdown report (headings, paragraphs, markdown tables) to DOCX."""

from __future__ import annotations

import argparse
from pathlib import Path
from typing import List

from docx import Document


def is_table_sep(line: str) -> bool:
    s = line.strip()
    if not s.startswith("|") or not s.endswith("|"):
        return False
    core = s.strip("|").replace("-", "").replace(":", "").replace(" ", "")
    return core == ""


def parse_table_block(lines: List[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines):
        line = lines[i].rstrip("\n")
        if not line.strip().startswith("|"):
            break
        if is_table_sep(line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    for r_i, row in enumerate(rows):
        for c_i in range(ncols):
            txt = row[c_i] if c_i < len(row) else ""
            table.cell(r_i, c_i).text = txt


def convert(md_path: Path, docx_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8", errors="ignore")
    lines = text.splitlines()

    doc = Document()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph("")
            i += 1
            continue

        if stripped.startswith("|"):
            rows, new_i = parse_table_block(lines, i)
            add_table(doc, rows)
            i = new_i
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            i += 1
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
            i += 1
            continue

        # simple bullet support
        if stripped.startswith("1. ") or stripped.startswith("2. ") or stripped.startswith("3. ") or stripped.startswith("4. ") or stripped.startswith("5. "):
            doc.add_paragraph(stripped, style="List Number")
            i += 1
            continue

        doc.add_paragraph(line)
        i += 1

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)


def main() -> None:
    p = argparse.ArgumentParser(description="Convert markdown report to DOCX")
    p.add_argument("--input", default="output_country_discourse_raw_analysis_strict/report_strict_method.md")
    p.add_argument("--output", default="output_country_discourse_raw_analysis_strict/report_strict_method.docx")
    args = p.parse_args()

    md_path = Path(args.input)
    if not md_path.exists():
        raise FileNotFoundError(f"Markdown report not found: {md_path}")

    docx_path = Path(args.output)
    convert(md_path, docx_path)
    print(f"DOCX generated: {docx_path.resolve()}")


if __name__ == "__main__":
    main()
